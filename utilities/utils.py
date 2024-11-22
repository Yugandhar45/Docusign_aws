import base64
from datetime import datetime
from io import BytesIO
from pathlib import Path

import pypdf
from docx import Document
from docx.shared import Inches, RGBColor, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.wait import WebDriverWait
from selenium.webdriver.common.by import By
from testData import constants as constants
import os
import time
from PIL import Image, ImageDraw, ImageFont
import logging
from logging.handlers import RotatingFileHandler
import random
import string
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Initialize logger and file handler outside the method
log_formatter = logging.Formatter('%(asctime)s %(levelname)s %(message)s ', datefmt='%Y-%m-%d %H:%M:%S')


class Util_Test:
    folder_path = constants.screenshots_path
    logs_folder_path = constants.custom_logs_path
    test_name = constants.test_name
    testCaseNum = 1


    def __init__(self, driver):
        self.driver = driver

        # Locators:
        self.profile_button = "//button[@data-qa='header-profile-menu-button']"
        self.logoff_button = "button[data-qa='header-logoff-button']"
        self.settings_tab = "//button[@data-qa='header-RADMIN-tab-button']"
        self.regional_settings_field = "//button[@data-qa='nav_link_authenticated.regional-settings']"
        self.regional_settings = "//button[@data-qa='regional-settings']"
        self.allow_user_set_time_zone_format = \
            "//button[@data-qa='regional_settings_allow_user_to_set_time_zone_format_cb']"
        self.dat_time_format_dropdown = "//select[@data-qa='date-time-format-dropdown']"

    def logout(self):
        WebDriverWait(self.driver, 60).until(EC.visibility_of_element_located((
            By.XPATH, self.profile_button))).click()
        WebDriverWait(self.driver, 60).until(EC.element_to_be_clickable((
            By.CSS_SELECTOR, self.logoff_button))).click()
        loginpage = WebDriverWait(self.driver, 45).until(EC.visibility_of_element_located((
            By.XPATH, "//span [contains(text(),'Log in to Docusign')]"))).is_displayed()
        assert loginpage

    @staticmethod
    def scroll_page(self):
        self.driver.find_element('css selector', "button[data-qa='tutorial-got-it']").click()
        time.sleep(5)
        self.driver.execute.script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(5)

    def switch_to_window(self):
        parentWindow = self.driver.current_window_handle
        main_window = self.driver.window_handles
        for handle in self.driver.window_handles:
            if handle != main_window:
                popup = handle
                self.driver.switch_to.window(popup)
                yield
        self.driver.switch_to.window(parentWindow)

    @staticmethod
    def validate_pdf_data(filecontents, first_page=False):
        file = open(filecontents[0], "rb")
        reader = pypdf.PdfReader(file)
        data = ""
        for page in reader.pages:
            data += page.extract_text()
            print(data)
        if first_page:
            pass
        else:
            assert filecontents[2] in data
        assert filecontents[1] in data

    def create_directory(self, test_name, root_directory=None):
        timestamp = datetime.now().strftime("%Y-%m-%d-%H-%M-%S")
        Util_Test.test_name = test_name
        folder_name = test_name + '_' + timestamp
        # global folder_path
        if root_directory is None:
            root_directory = os.getcwd()
            Util_Test.folder_path = os.path.join(root_directory, 'screenshots', folder_name)
        os.makedirs(Util_Test.folder_path)

    # @staticmethod
    # def create_directory_for_customlogs(self, test_name, root_directory=None):
    #     timestamp = datetime.now().strftime("%Y-%m-%d-%H-%M-%S")
    #     folder_name = test_name + '_' + timestamp
    #     # global folder_path
    #     if root_directory is None:
    #         root_directory = os.getcwd()
    #         Util_Test.logs_folder_path = os.path.join(root_directory, 'executionlogs', folder_name)
    #     os.makedirs(Util_Test.logs_folder_path)

    def getscreenshot(self, fileName):
        screenshot = self.driver.get_screenshot_as_png()
        image = Image.open(BytesIO(screenshot))
        current_datetime = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        draw = ImageDraw.Draw(image)
        font_size = 16
        font = ImageFont.truetype("arial.ttf", font_size)
        text_padding = 5
        bbox = draw.textbbox((0, 0), current_datetime, font=font)
        text_width = bbox[2] - bbox[0]
        text_height = bbox[3] - bbox[1]
        image_width, image_height = image.size
        text_position = (text_padding, image_height - text_height - text_padding)
        draw.text(text_position, current_datetime, fill="black", font=font)
        filepath = os.path.abspath(Util_Test.folder_path) + '/' + fileName
        image.save(filepath)
        return filepath

    @staticmethod
    def password_encrypt(*args):
        for value_to_encrypt in args:
            password = value_to_encrypt
            encrypt_string = base64.b64encode(password.encode("utf-8"))
            print(value_to_encrypt + ":", encrypt_string)

    @staticmethod
    def password_decrypt(value_to_decrypt):
        password = value_to_decrypt
        decrypt_string = base64.decodebytes(password).decode("utf-8")
        return decrypt_string

    # @staticmethod
    # def speak(text):
    #     engine = pyttsx3.init()
    #     engine.say(text)
    #     engine.runAndWait()

    def execute_script_with_banner(self, text, apply_fixed_position=True, isBanner=False):
        # Display Banner
        if isBanner:
            text_for_banner = f'<b>{text}</b>'
            script = f''' 
                        var banner = document.createElement('div'); 
                        banner.style.position = '{'fixed' if apply_fixed_position else 'relative'}'; 
                        banner.style.top = '0'; 
                        banner.style.left = '0'; 
                        banner.style.right = '0'; 
                        banner.style.margin = 'auto'; 
                        banner.style.width = '700px'; 
                        banner.style.height = '50px'; 
                        banner.style.backgroundColor = 'yellow'; 
                        banner.style.color = 'black'; 
                        banner.style.textAlign = 'center'; 
                        banner.style.lineHeight = '50px'; 
                        banner.style.zIndex = '9999'; 
                        banner.style.pointerEvents = 'none'; 
                        banner.innerHTML = '{text_for_banner}'; 
                        document.body.appendChild(banner); 
                    '''
            self.driver.execute_script(script)
            # Speak
            # self.speak(text)

    @staticmethod
    def initialize_logger(test_case):
        log_folder = "executionlogs"
        if not os.path.exists(log_folder):
            os.makedirs(log_folder)
        log_file = os.path.join(log_folder, f"{test_case}_execution.log")

        file_handler = RotatingFileHandler(log_file, mode='w')
        file_handler.setFormatter(log_formatter)
        file_handler.setLevel(logging.INFO)

        logger = logging.getLogger(test_case)
        logger.setLevel(logging.INFO)
        logger.addHandler(file_handler)

        return logger

    @staticmethod
    def write_custom_logs(logger, line):
        logger.info(line)

    @staticmethod
    def get_random_code():
        code = ''.join(random.choices(string.ascii_uppercase + string.digits, k=6))
        return code

    @staticmethod
    def add_test_name_to_doc(testcasename):
        doc_path = Util_Test.document_path()
        doc = Document(doc_path)
        testcasenum = f"TestCase{Util_Test.testCaseNum} :"
        testname = testcasenum+testcasename
        #doc.add_paragraph()
        doc.add_heading(testname, level=1)
        doc.save(doc_path)

    @staticmethod
    def add_failed_message_doc(testname):
        doc_path = Util_Test.document_path()
        doc = Document(doc_path)
        image_folder = Util_Test.folder_path
        for image_name in sorted(os.listdir(image_folder)):
            if image_name.startswith('Failed_Screenshot.png'):
                # Full path to the image
                image_path = os.path.join(image_folder, image_name)
                image_name = image_name.split(".")
                # Add a paragraph with the image name (optional)
                doc.add_paragraph(image_name[0])
                # Append the image to the document
                doc.add_picture(image_path, width=Inches(7.0), height=Inches(3.8))
                break
        paragraph = doc.add_paragraph()
        run1 = paragraph.add_run(testname + '----' + ' Script Failed')
        run1.font.color.rgb = RGBColor(255, 0, 0)  # Red color
        doc.save(doc_path)

    @staticmethod
    def add_screenshots_to_doc():

        doc_path = Util_Test.document_path()
        doc = Document(doc_path)
        if doc:
            print("Document loaded successfully.")
        image_folder = Util_Test.folder_path
        for image_name in sorted(os.listdir(image_folder)):
            if image_name.endswith('.png'):
                # Full path to the image
                image_path = os.path.join(image_folder, image_name)
                # Add a paragraph with the image name (optional)
                image_name = image_name.split(".png")
                doc.add_paragraph(image_name[0])
                # Append the image to the document
                doc.add_picture(image_path, width=Inches(7.0), height=Inches(3.8))  # Adjust the width as needed
        # Save the document with the newly appended images
        doc.save(doc_path)

    @staticmethod
    def document_path():
        summary_report = constants.summary_report_path
        root_directory = os.getcwd()
        relative_path = Path(summary_report)
        doc_path = root_directory / relative_path
        file_path = Path(doc_path)
        summary_report_path = str(file_path)
        return summary_report_path

    @staticmethod
    def create_document():
        doc = Document()
        summary_report_path = Util_Test.document_path()
        for x in range(5):
            doc.add_paragraph()
        # Add title to the cover page
        title = doc.add_paragraph()
        title_run = title.add_run("DocuSign Automation")
        title_run.font.size = Pt(28)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Add subtitle to the cover page
        subtitle = doc.add_paragraph()
        subtitle_run = subtitle.add_run("Test Summary Report")
        subtitle_run.font.size = Pt(22)
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Add date
        date = doc.add_paragraph()
        date_run = date.add_run(datetime.now().strftime("%B %d, %Y"))
        date_run.font.size = Pt(16)
        date.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # Add a header
        header = doc.sections[0].header
        header_paragraph = header.paragraphs[0]
        # Insert the logo image in the header
        logo_run = header_paragraph.add_run()
        logo_path = os.path.join(os.path.dirname("tests"), 'reports', constants.logo_path)
        logo_run.add_picture(logo_path, width=Inches(1.5))
        # Add a footer
        footer = doc.sections[0].footer
        # Add a paragraph to the footer for the page number
        footer_paragraph = footer.paragraphs[0]
        footer_paragraph.text = "Page "  # Static text before the page number

        # Create the "PAGE" field element
        page_number_run = footer_paragraph.add_run()
        fld_char = OxmlElement("w:fldChar")  # Start of field
        fld_char.set(qn("w:fldCharType"), "begin")
        instr_text = OxmlElement("w:instrText")  # Field instructions
        instr_text.text = "PAGE"
        fld_char_separate = OxmlElement("w:fldChar")  #Separator between field code and result
        fld_char_end = OxmlElement("w:fldChar")  # End of field
        fld_char_end.set(qn("w:fldCharType"), "end")
        page_number_run.font.color.rgb = None

        # Add field elements in the correct order
        page_number_run._r.append(fld_char)
        page_number_run._r.append(instr_text)
        page_number_run._r.append(fld_char_separate)
        page_number_run._r.append(fld_char_end)
        doc.add_page_break()
        doc.save(summary_report_path)


