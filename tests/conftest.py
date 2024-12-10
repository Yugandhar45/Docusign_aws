
import base64

from docx import Document
from selenium import webdriver
from selenium.common import NoSuchWindowException
from selenium.webdriver.chrome.service import Service as ChromeService
from selenium.webdriver.edge.service import Service as EdgeService
from selenium.webdriver.firefox.service import Service as FirefoxService
from webdriver_manager.chrome import ChromeDriverManager
from webdriver_manager.microsoft import EdgeChromiumDriverManager
from webdriver_manager.firefox import GeckoDriverManager
from utilities.utils import Util_Test
from datetime import datetime
from testData import constants
import pytest
import os
import pytz
from py.xml import html



def pytest_addoption(parser):
    parser.addoption("--browser", action="store", default="edge")

@pytest.fixture(scope="session", autouse=True)
def suite_setup():
    Util_Test.create_document()

@pytest.fixture(scope='class')
def test_setup(request):
    driver = None
    browser = request.config.getoption("--browser")
    if browser == "chrome":
        download_path = os.path.abspath(constants.download_path)
        options = webdriver.ChromeOptions()
        options.add_argument("disable-features=DownloadUI")
        options.add_experimental_option("prefs", {
            "download.default_directory": download_path,
            "plugins.plugins_disabled": ["Chrome PDF Viewer"],
            "download.prompt_for_download": False,
            "download.directory_upgrade": True,
            "safebrowsing.enabled": True
        })
        driver = webdriver.Chrome(service=ChromeService(ChromeDriverManager().install()), options=options)
    elif browser == "firefox":
        download_path = os.path.abspath(constants.download_path)
        options = webdriver.FirefoxOptions()
        options.set_preference("browser.download.folderList", 2)
        options.set_preference("browser.download.manager.showWhenStarting", False)
        options.set_preference("signon.management.page.os-authKeystore", False)
        options.set_preference("browser.download.dir", download_path)
        options.set_preference("browser.helperApps.neverAsk.saveToDisk", "application/pdf")
        driver = webdriver.Firefox(service=FirefoxService(GeckoDriverManager().install()), options=options)
    elif browser == "edge":
        download_path = os.path.abspath(constants.download_path)
        edge_options = webdriver.EdgeOptions()
        edge_options.add_argument("--inprivate")
        edge_options.add_experimental_option('prefs', {
            'download.default_directory': download_path,
            'download.prompt_for_download': False,
            'download.directory_upgrade': True,
            'safebrowsing.enabled': True
        })
        driver = webdriver.Edge(service=EdgeService(EdgeChromiumDriverManager().install()), options=edge_options)
    driver.implicitly_wait(2)
    driver.maximize_window()
    driver.delete_all_cookies()
    request.cls.driver = driver
    yield
    Util_Test.testCaseNum = Util_Test.testCaseNum + 1
    driver.quit()


# HTML Reports
def pytest_html_report_title(report):
    report.title = "Docusign Automation Test Report"


@pytest.hookimpl(hookwrapper=True)
def pytest_runtest_makereport(item):
    pytest_html = item.config.pluginmanager.getplugin("html")
    outcome = yield
    report = outcome.get_result()
    extra = getattr(report, 'extra', [])
    if report.when == 'call':
        xfail = hasattr(report, 'wasxfail')
        if (report.skipped and xfail) or (report.failed and not xfail):
            driver = getattr(item.cls, 'driver', None)
            if driver:
                util_test = Util_Test(driver)
                screenshots_dir = 'screenshots'
                if not os.path.exists(screenshots_dir):
                    os.makedirs(screenshots_dir)
                try:
                    screenshot_path = util_test.getscreenshot("Failed_Screenshot.png")
                    Util_Test.add_failed_message_doc(Util_Test.test_name)
                    if os.path.exists(screenshot_path):
                        with open(screenshot_path, "rb") as image_file:
                            encoded_string = base64.b64encode(image_file.read()).decode("utf-8")
                            extra_html = (
                                f'<div style="float: right; margin-left: 20px;">'
                                f'<img src="data:image/png;base64,{encoded_string}" alt="screenshot" '
                                f'style="width:200px; height:150px;" onclick="window.open(this.src)" align="right"/></div>'
                            )
                            extra.append(pytest_html.extras.html(extra_html))
                except NoSuchWindowException:
                    extra_html = (
                        '<div style="float: right; margin-left: 20px;">'
                        '<p>Screenshot could not be captured as the window was closed.</p>'
                        '</div>'
                    )
                    extra.append(pytest_html.extras.html(extra_html))
    report.extra = extra


# It is the Hook to add the logo in the Report
@pytest.hookimpl(tryfirst=True)
def pytest_html_results_summary(prefix):
    logo_path = os.path.join(os.path.dirname(__file__), 'reports', constants.logo_path)
    if os.path.exists(logo_path):
        prefix.extend([html.div(
            html.img(src=logo_path, alt="Logo", style="position:absolute;top:20px;right:10px;padding:40px", height="50",
                     width="100")
        )])


# It is the hook for adding environment info to html reports
def pytest_configure(config):
    config._metadata['Project Name'] = 'DocuSign'
    config._metadata['Run User'] = os.environ.get('TriggeringUser', 'Unknown')
    config._metadata['UTC Time'] = datetime.now(pytz.UTC)


# It is Hook for delete/modify environment info to HTML report
def pytest_metadata(metadata):
    metadata.pop("Packages", None)
    metadata.pop("Plugins", None)
    metadata.pop("Python", None)


def pytest_terminal_summary(terminalreporter):
    # Getting the test results
    passed_tests = terminalreporter.stats.get('passed', [])
    failed_tests = terminalreporter.stats.get('failed', [])
    skipped_tests = terminalreporter.stats.get('skipped', [])

    passed = len(passed_tests)
    failed = len(failed_tests)
    skipped = len(skipped_tests)

    # Creating a Word document
    doc_path = Util_Test.document_path()
    doc = Document(doc_path)
    # Writing passed tests
    doc.add_heading('Passed Tests', level=1)
    if passed_tests:
        i = 1
        for test in passed_tests:
            testname = test.nodeid.split('::')
            doc.add_paragraph("Test Script - {} : {} ".format(i, testname[2]))
            i += 1
    else:
        doc.add_paragraph("No tests passed.")

    # Writing failed tests
    if failed_tests:
        doc.add_heading('Failed Tests', level=1)
        i = 1
        for test in failed_tests:
            testname = test.nodeid.split('::')
            doc.add_paragraph("Test Script - {} : {} ".format(i, testname[2]))
            i += 1

    # Writing skipped tests
    if skipped_tests:
        doc.add_heading('Skipped Tests', level=1)
        i = 1
        for test in skipped_tests:
            testname = test.nodeid.split('::')
            doc.add_paragraph("Test Script - {} : {} ".format(i, testname[2]))
            i += 1
    # Summary of results
    doc.add_heading('Summary', level=1)
    doc.add_paragraph(f"Total Passed: {passed}")
    doc.add_paragraph(f"Total Failed: {failed}")
    doc.add_paragraph(f"Total Skipped: {skipped}")
    doc.save(doc_path)
    print(f"\nPassed: {passed}, Failed: {failed}, Skipped: {skipped}")
